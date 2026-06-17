from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from pathlib import Path
import csv


OUT = Path("博格华纳项目开发管理模块_需求规格说明书_V1.0.docx")


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def p(doc, text="", style=None, bold_prefix=None):
    para = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = para.add_run(bold_prefix)
        set_run_font(r, bold=True)
        r2 = para.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        run = para.add_run(text)
        set_run_font(run)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    set_run_font(run)
    return para


def numbered(doc, text):
    para = doc.add_paragraph(style="List Number")
    run = para.add_run(text)
    set_run_font(run)
    return para


def table(doc, headers, rows, widths=None, header_fill="F2F4F7"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.autofit = False
    set_table_width(tbl)
    set_table_borders(tbl)
    set_cell_margins(tbl)
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_shading(c, header_fill)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(c, widths[i])
        r = c.paragraphs[0].add_run(h)
        set_run_font(r, bold=True, size=9.5, color="1F2937")
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cells[i], widths[i])
            run = cells[i].paragraphs[0].add_run(str(value))
            set_run_font(run, size=9)
    doc.add_paragraph()
    return tbl


def callout(doc, title, body, fill="F4F6F9"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(tbl)
    set_table_borders(tbl, color="C9D3E1")
    set_cell_margins(tbl, top=120, bottom=120, start=160, end=160)
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, fill)
    pr = cell.paragraphs[0]
    r = pr.add_run(title)
    set_run_font(r, bold=True, color="1F4D78")
    pr.add_run("\n")
    r2 = pr.add_run(body)
    set_run_font(r2, size=10)
    doc.add_paragraph()


def add_title(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("博格华纳-项目开发管理模块\n需求规格说明书")
    set_run_font(run, size=22, bold=True, color="0B2545")
    para.paragraph_format.space_after = Pt(10)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("版本：V1.0    日期：2026-06-16    用途：客户范围确认 / 开发实现依据")
    set_run_font(r, size=10.5, color="555555")
    doc.add_paragraph()
    callout(
        doc,
        "文档定位",
        "本文基于当前原型页面、AGENTS.md 项目需求上下文、权限设计说明、Gate Checklist 拆分数据及 BOM/APQP 调整方案整理。用于确认项目开发管理模块范围，并作为后续前后端开发、测试和验收的需求依据。"
    )


def get_gate_summary_rows():
    path = Path("data/gate_tasks_summary_by_ui_stage.csv")
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return [[row["uiStage"], row["total"]] for row in csv.DictReader(f)]


def main():
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. 文档说明", level=1)
    table(doc, ["项目", "说明"], [
        ["文档目的", "明确项目开发管理模块的业务流程、角色权限、页面需求、交互规则、数据字段和验收边界。"],
        ["阅读对象", "客户业务负责人、项目经理、产品经理、UI/前端开发、后端开发、测试人员、实施顾问。"],
        ["适用范围", "项目创建、CKO、PKO、Gate1-Gate6、Hand Over、BOM、APQP、待办、审批、权限、看板、基础配置。"],
        ["原型依据", "项目开发管理.html、index-project-record-bom-supply-apqp.html、待办任务中心-关键节点方案.html、子任务库管理.html、子零件开发任务详情_调整版.html、项目数据统计看板.html。"],
        ["数据依据", "data/gate_tasks_split.ui.csv、data/gate_tasks_summary_by_ui_stage.csv、权限设计说明.md、AGENTS.md。"],
    ], [2100, 7200])

    doc.add_heading("2. 项目背景与建设目标", level=1)
    p(doc, "本模块面向制造业新产品开发项目全生命周期管理，以门径管理模式为核心框架。项目从销售创建开发项目开始，依次流转 CKO、PKO、Gate1、Gate2 至 Gate6、Hand Over，并由销售、PM、AE、SCM、AQE、SQD 等跨职能角色协同执行。")
    p(doc, "建设目标包括：统一项目主数据入口，固化 CKO/PKO/Gate/Hand Over 阶段流转；将 Gate Checklist 转化为可分配、可跟踪、可审核的任务实例；建立 BOM、供应方案、APQP 与子零件开发任务之间的联动关系；通过待办中心和看板提升跨部门协同效率。")
    callout(doc, "核心确认口径", "项目 Gate 推进不只看项目级任务，还必须校验 BOM 状态、L1 APQP、子零件 APQP 和阻塞型供应商开发任务。当前 Gate 下全部必做任务完成后，PM 才可提交阶段审核。")

    doc.add_heading("3. 总体业务流程", level=1)
    table(doc, ["阶段", "主责角色", "核心动作", "流转条件"], [
        ["CKO", "销售 Sales", "创建开发项目，完成系统默认 CKO 任务。", "CKO 任务完成后提交至 PKO。"],
        ["PKO", "PM、AE、SCM", "PM 发起启动会，录入关键节点，选择标准/简化流程；AE 创建 BOM，SCM 确认 BOM。", "SCM 确认并锁定 BOM 后进入 Gate1。"],
        ["Gate1", "PM、AQE、SQD", "PM 基于 L1 BOM 创建总成 APQP；AQE 拆解子零件 APQP；SQD 接收供应商开发任务。", "Gate1 当前任务、BOM、APQP 准入检查通过后进入 Gate2。"],
        ["Gate2-Gate6", "PM 主导，各职能执行", "PM 分配阶段任务，各角色填报完成情况，审核小组完成 Gate Review。", "当前 Gate 全部必做任务及阻塞型子零件任务完成，审核通过后进入下一 Gate。"],
        ["Hand Over", "PM", "登记交接信息、交接团队、资料和未关闭事项。", "交接信息完成后项目结束。"],
    ], [1200, 1600, 4100, 2500])

    doc.add_heading("3.1 标准流程与简化流程", level=2)
    table(doc, ["流程类型", "适用范围", "阶段范围", "规则"], [
        ["标准流程", "VT 系统：Phaser、VFS、CB、CCV；TD 系统：Chain、TSC、Timing、Tensioner、OP、CC Guide、Arm Guide。", "Gate1 至 Gate6", "生成并跟踪 Gate1-Gate6 任务；Gate4 对应 BW 内部 OTS，Gate5 对应 BW 内部 PPAP，Gate6 对应 BW 内部 SOP。"],
        ["简化流程", "PP/Chain 内部子零件、Motor Chain 场景。", "Gate1 至 Gate3", "只生成 Gate1-Gate3 任务；Gate3 完成后的下一阶段是否直接进入 Hand Over 需客户确认。"],
    ], [1400, 3500, 1600, 2800])

    doc.add_heading("3.2 关键节点映射", level=2)
    table(doc, ["节点类型", "字段", "映射/用途"], [
        ["客户节点", "客户 OTS、客户 PPAP、客户 SOP", "记录客户承诺时间，用于管理层看板、预警和内部计划校验。"],
        ["BW 内部节点", "BW OTS、BW PPAP、BW SOP", "分别对应 Gate4、Gate5、Gate6。"],
        ["Gate 节点", "PKO、Gate1、Gate2、Gate3/OTS、Gate4/PPAP、Gate5/SOP", "项目计划页用于维护 Gate 基线，校验阶段日期顺序及是否晚于客户节点。"],
    ], [1600, 3000, 4700])

    doc.add_heading("4. 用户角色与权限", level=1)
    table(doc, ["角色", "主要职责", "典型页面/操作"], [
        ["销售 Sales", "创建开发项目，维护客户输入，完成 CKO 任务。", "项目列表、新增项目、CKO 任务填报。"],
        ["PM", "管理项目成员、发起 PKO、维护计划、选择流程、分配任务、推进 Gate、发起审核、Hand Over。", "项目详情、任务管理、Gate 提交、项目计划、团队成员、交接登记。"],
        ["AE", "创建并维护产品 BOM 初稿，提交 SCM 确认。", "BOM 编辑、BOM 提交、工程任务填报。"],
        ["SCM", "确认 BOM 供应链信息、供应方案和供应商准备状态，锁定 BOM。", "确认 BOM、供应方案维护、BOM 变更确认。"],
        ["AQE", "默认 L1 APQP 负责人，管理总成 APQP，拆解子零件任务。", "APQP 创建、子零件 APQP 分配、质量任务。"],
        ["SQD", "默认子零件 APQP 负责人，对接供应商，跟踪子零件开发节点。", "子零件任务详情、供应商节点同步、Gate 任务填报。"],
        ["审核小组", "对阶段审核或任务完成申请给出结论。", "待办中心审核、Gate Review。"],
        ["超级管理员", "维护阶段、任务模板、角色、权限、异常兜底。", "子任务库管理、配置管理、异常处理。"],
    ], [1300, 4700, 3300])
    p(doc, "权限模型建议采用“项目成员可见 + 具体责任人可操作 + 责任部门作为默认分配范围”。后端必须做权限校验，前端按钮隐藏仅作为体验优化。")
    table(doc, ["操作", "PM", "任务责任人", "同部门非责任人", "APQP 负责人", "审核小组", "超级管理员"], [
        ["查看项目", "是", "项目成员可看", "项目成员可看", "项目成员可看", "参与审核范围可看", "是"],
        ["编辑项目主数据", "是", "否", "否", "否", "否", "是"],
        ["分配/改派 Gate 任务", "是", "否，除非授权", "否", "管理范围内可分配", "否", "是"],
        ["填报/提交任务", "可代办需留痕", "是", "否", "管理范围内可操作", "否", "可兜底"],
        ["发起 Gate 审核", "是", "否", "否", "否", "否", "是"],
        ["Gate 审核", "按配置", "否", "否", "按配置", "是", "是"],
        ["Hand Over 登记", "是", "否", "否", "否", "否", "是"],
    ], [1550, 1200, 1450, 1550, 1450, 1300, 1300])

    doc.add_heading("5. 功能模块总览", level=1)
    table(doc, ["模块", "页面/功能", "说明"], [
        ["项目管理", "项目列表、新增项目、项目详情、项目状态控制", "项目主入口，覆盖项目创建、查询、阶段状态和项目基础信息。"],
        ["项目开发任务", "阶段时间轴、任务清单、任务填报、审核记录、任务管理", "按 CKO/PKO/Gate/Hand Over 展示阶段任务，支持任务筛选、填报、提交审核。"],
        ["BOM 管理", "生产 BOM 列表、BOM 编辑、SCM 确认、历史版本、受控导出、变更影响", "管理产品 BOM 层级、供应方案、变更状态和 APQP 生成策略。"],
        ["APQP 管理", "一键创建 APQP、APQP 任务列表、子零件 APQP", "基于锁定 BOM 生成 L1 APQP 和子零件 APQP。"],
        ["待办中心", "角色切换、待处理/已完成、任务处理、审核处理", "按角色和任务来源汇总待办，支持处理、完成、审核。"],
        ["客户要求", "客户需求列表、添加/编辑/删除", "记录项目需求输入和变更。"],
        ["问题与 8D", "问题清单、质量问题 8D", "支持项目问题登记、状态跟踪和质量问题管理。"],
        ["团队成员", "项目团队成员、审核小组、成员选择", "维护项目成员和审核人员。"],
        ["资料库", "项目开发任务附件、手动上传资料", "统一沉淀交付物、附件和项目资料。"],
        ["配置与看板", "子任务库管理、项目数据统计看板", "维护阶段/任务模板，输出项目阶段、节点、风险和负载统计。"],
    ], [1500, 2800, 5000])

    doc.add_heading("6. 页面与模块需求说明", level=1)
    doc.add_heading("6.1 项目列表页", level=2)
    p(doc, "使用角色：销售、PM、项目成员、超级管理员。页面用于查询项目、查看项目状态和进入项目详情，销售可新增项目。")
    table(doc, ["区域", "需求说明"], [
        ["筛选区", "项目编号、项目名称、客户品牌、状态；支持查询、重置和自定义筛选条件。"],
        ["列表区", "支持勾选、横向固定列、分页；左侧固定项目编号/名称，右侧固定阶段/状态/操作。"],
        ["操作区", "新增项目、查看详情；后续可扩展批量导出或批量状态操作。"],
    ], [1800, 7500])
    table(doc, ["字段", "类型", "说明"], [
        ["项目编号", "文本", "项目唯一编号，如 DMSNIN25001。"],
        ["项目名称", "文本", "客户项目或产品开发项目名称。"],
        ["销售责任单元", "文本/枚举", "如 NINGBO VT、NINGBO TD。"],
        ["客户", "枚举", "客户品牌。"],
        ["系统", "枚举/自动带出", "由所选产品自动归类，如 VVT 系统、TD、PP。"],
        ["产品", "多选枚举", "Phaser、VFS、CB、OCV、Chain、TSC、Timing Tensioner、OP tensioner、Sprocket、CC Guide、Arm Guide、PP(Chain内部子零件)、Motor Chain。"],
        ["项目经理", "人员", "PM 项目经理。"],
        ["等级", "枚举", "A/B/C。"],
        ["物流计划类型", "枚举", "Booked、未预定、LRP+。"],
        ["峰值订单量/预计营收/配套需求量", "数字", "销售与市场数据。"],
        ["定点季度/定点日期/SOP", "季度/日期", "项目商业节点。"],
        ["阶段/状态", "枚举", "阶段：CKO、PKO、Gate1-Gate6、Hand Over；状态：进行中、审批中、已拒绝、已通过、带条件通过等。"],
    ], [2000, 1500, 5800])

    doc.add_heading("6.2 新增项目抽屉页", level=2)
    p(doc, "销售人员通过抽屉式表单创建项目。确认创建后弹出“项目已创建”提示，可选择现在填写 CKO 阶段任务或稍后填写。")
    table(doc, ["分区", "字段/功能", "规则"], [
        ["基本信息", "销售责任单元、客户、项目、产品、系统、项目经理", "销售责任单元、客户、项目、产品、项目经理必填；系统根据产品选择自动带出。"],
        ["销售 & 前瞻规划", "峰值订单量、定价、预计营收、配套需求量", "数字字段需校验格式；可为空但应支持后续补充。"],
        ["采购与排程", "定点季度、定点日期", "日期选择器支持清空和确定。"],
        ["评价与分级", "SOP、等级、物流计划类型、RYG 评分", "RYG 评分覆盖开发跟进、财务、制造、价格成本、产品、品质、关系、供应商、时间节点。"],
        ["备注", "备注", "长文本。"],
        ["底部操作", "取消、保存草稿、确定", "保存草稿不进入正式流转；确定后生成项目并进入 CKO。"],
    ], [1600, 3900, 3800])

    doc.add_heading("6.3 项目详情页", level=2)
    p(doc, "项目详情页是项目执行主工作台，包含项目记录、项目开发任务、生产 BOM、客户要求、项目计划、问题清单、质量问题 8D、团队成员、资料库等页签。")
    table(doc, ["页签", "页面描述", "主要功能点"], [
        ["项目记录", "展示当前提醒、项目状态摘要和项目动态。", "按全部/阶段/BOM/任务/审批/成员/状态过滤项目动态；记录项目状态变更、BOM 提交、APQP 创建等事件。"],
        ["项目开发任务", "按阶段时间轴展示 CKO、PKO、Gate1-Gate6、Hand Over。", "查看阶段说明；按 ID/Title、部门、状态筛选；任务填报；查看审核记录；进入任务管理；提交阶段审核。"],
        ["生产 BOM", "展示当前及历史 BOM、供应方案、APQP 策略。", "查看详情、预览 BOM、BOM 变更影响、导出 BOM、APQP 任务、一键创建 APQP、编辑 BOM（AE）、确认 BOM（SCM）。"],
        ["客户要求", "维护客户需求输入。", "添加客户需求、详情、编辑、删除。"],
        ["项目计划", "维护客户、BW 内部和 Gate 时间节点。", "保存计划、恢复默认、校验日期顺序与客户节点约束。"],
        ["问题清单", "管理项目问题。", "批量导入、导出、添加问题、筛选待解决/已解决、查询/重置、详情/编辑/删除。"],
        ["质量问题 8D", "管理质量问题闭环。", "按关键信息查询；查看/管理 8D 问题。"],
        ["团队成员", "维护项目团队和审核小组。", "添加成员、查询、移除成员、设置审核小组。"],
        ["资料库", "统一管理附件资料。", "上传附件；按全部资料、项目开发任务附件、手动上传资料筛选。"],
    ], [1400, 3600, 4300])

    doc.add_heading("6.4 项目计划页", level=2)
    table(doc, ["分区", "字段", "交互/逻辑"], [
        ["客户版时间节点", "Quote Due、Nomination、Prototype、客户 OTS、客户 PPAP、客户 SOP", "用户手动维护，作为内部计划约束和看板预警依据。"],
        ["BW 内部版时间节点", "PM Kick off、AR Approval、Tooling kick off、Prototype、BW OTS、BW PPAP、BW SOP", "BW OTS/PPAP/SOP 分别映射 Gate4/5/6。"],
        ["Gate 时间节点", "PKO、Gate1、Gate2、Gate3/OTS、Gate4/PPAP、Gate5/SOP", "校验顺序：PKO 不晚于 Gate1，Gate1 不晚于 Gate2，依次类推；BW 内部计划不得晚于客户要求节点。"],
        ["展示", "客户/BW 双时间轴、Gate 计划基线", "用于对比客户承诺与内部计划。"],
    ], [1900, 3600, 3800])

    doc.add_heading("6.5 项目开发任务页", level=2)
    table(doc, ["项目", "说明"], [
        ["使用角色", "PM、任务责任人、APQP 负责人、审核小组、超级管理员。"],
        ["任务来源", "Gate Checklist UI 数据、系统默认任务、PM 手动新增、子任务库配置、BOM/APQP 自动生成任务。"],
        ["列表字段", "编号、K-PAC ID、K-PAC Title、Description、负责人、负责部门、任务类型、起止时间、Document or Evidence、状态、操作。"],
        ["筛选条件", "ID / Title、部门、状态（全部、待完成、已逾期、已完成）。"],
        ["功能点", "阶段切换、展开/收起、任务填报、附件上传、提交完成、查看审核记录、进入任务管理、阶段提交审核。"],
        ["任务填报", "显示任务标题、描述、交付物要求；支持上传附件和填写说明；提交后进入待审核或已完成状态，具体按任务配置。"],
        ["阶段提交", "PM 点击提交时进行准入检查；若存在未完成项目任务、未处理 BOM 变更、阻塞型子零件 APQP 未完成，则禁止提交并列出原因。"],
    ], [1800, 7500])

    doc.add_heading("6.6 BOM 管理", level=2)
    p(doc, "项目 BOM 是子零件开发任务和 APQP 任务的主要来源。BOM 未经 SCM 确认锁定前，不允许生成正式 L1 APQP 和子零件 APQP。")
    table(doc, ["页面/弹窗", "使用角色", "功能点"], [
        ["生产 BOM 列表", "PM、AE、SCM、AQE、SQD", "查看当前 BOM、历史版本、状态、行数、BOM 变更数量；筛选关键信息、BOM 版本、状态。"],
        ["编辑 BOM（AE）", "AE", "添加产品 BOM、从零件库添加、手动添加子项、编辑行、删除行、复制/批量删除、保存草稿、提交 SCM。"],
        ["确认 BOM（SCM）", "SCM", "补充供应方案、确认供应链信息、识别变更影响、确认并锁定 BOM。"],
        ["BOM 详情/历史版本", "项目成员只读", "按版本查看 BOM 明细，不影响当前编辑数据。"],
        ["BOM 变更影响", "PM、SCM、AE", "展示相对上一锁定版本的变化类型、是否阻塞 Gate、建议处理动作。"],
        ["受控导出", "PM、SCM", "导出 Excel、PDF 或 Excel+PDF；锁定版本可生成受控编码。"],
    ], [2200, 1700, 5400])
    table(doc, ["BOM 字段", "类型", "说明"], [
        ["序号/级数", "数字/层级", "支持 L1 总成、L2/L3/L4 子层级。"],
        ["零件/材料名称、产品应用、数量", "文本/数字", "描述 BOM 行基础信息。"],
        ["客户零件号、BW 零件号、图纸版本", "文本", "工程识别信息。"],
        ["样件 SAP 零件号、量产 SAP 零件号、原材料零件号", "文本", "SAP 和物料信息。"],
        ["原材料、重量 Piece、重量 Assembly", "文本/数字", "材料与重量信息。"],
        ["共用零件、借用项目、图纸状态", "枚举/文本", "共用件默认可不生成子零件 APQP；图纸状态如 R2/R3。"],
        ["原材料供应商、分供方", "文本/供应商", "供应链基础信息。"],
        ["一供、二供", "供应方案", "不作为普通文本，需纳入供应方案管理。"],
        ["供应策略", "枚举", "单一供应、双供备案、双供同步开发、二供后续导入。"],
        ["APQP 生成策略", "枚举", "生成 L1 APQP、生成子零件 APQP、不生成。"],
        ["是否阻塞 Gate、起始 Gate", "枚举", "用于 Gate 准入检查。"],
        ["变更状态、APQP 状态", "枚举", "记录 BOM 变更和任务生成结果。"],
        ["备注/操作", "文本/按钮", "查看、编辑、删除等。"],
    ], [2300, 1400, 5600])
    table(doc, ["BOM 状态", "说明", "规则"], [
        ["草稿", "AE 或 SCM 编辑中。", "不影响当前已生效任务。"],
        ["待 SCM 确认", "AE 提交后等待 SCM 确认。", "不生成正式 APQP；可提醒 SCM 处理。"],
        ["锁定", "SCM 确认并锁定的当前有效版本。", "可作为 APQP 和子零件开发任务生成依据。"],
        ["历史版本", "每次锁定形成的版本快照。", "不可物理覆盖，支持追溯和只读查看。"],
    ], [1600, 3200, 4500])

    doc.add_heading("6.7 供应方案与 BOM 变更规则", level=2)
    table(doc, ["供应策略", "APQP 生成", "是否阻塞 Gate"], [
        ["单一供应", "默认生成一供子零件 APQP；L1 生成 L1 APQP。", "需开发时默认阻塞。"],
        ["双供备案", "二供仅备案，不生成 APQP。", "默认不阻塞。"],
        ["双供同步开发", "一供、二供均生成独立子零件 APQP。", "默认均阻塞。"],
        ["二供后续导入", "二供在指定 Gate 或指定条件后启动。", "启动前不阻塞，启动后按阻塞型任务处理。"],
    ], [2100, 4100, 3200])
    table(doc, ["变更类型", "处理原则", "默认 Gate 影响"], [
        ["新增零件", "判断是否需要开发；需要开发则生成子零件开发任务并确定起始 Gate。", "未判断或未生成任务时阻塞。"],
        ["删除零件", "不得物理删除已生成任务，应关闭或取消任务并保留历史。", "对应任务未关闭时阻塞。"],
        ["替换零件", "建立旧零件与新零件关系，判断任务继承、部分重做或完全重做。", "影响分析未完成时阻塞。"],
        ["一供变更", "视为高影响变更，触发供应商开发任务重评。", "默认阻塞。"],
        ["二供新增/删除/策略变化", "根据二供策略决定是否生成或关闭任务。", "备案不阻塞；同步开发或已启动任务未完成时阻塞。"],
        ["数量/层级/规格图纸变化", "记录变更，必要时通知 SCM 或重开相关 Gate 任务。", "默认不阻塞，若影响分析判定为关键影响则阻塞。"],
    ], [2100, 4700, 2500])

    doc.add_heading("6.8 APQP 管理", level=2)
    table(doc, ["项目", "需求说明"], [
        ["创建顺序", "项目 BOM 创建并经 SCM 确认锁定后，由 PM 创建 BOM 中 L1 层级 APQP；L1 层级必须创建。"],
        ["负责人", "创建 L1 APQP 时指定 L1 负责人，默认 AQE；创建子零件 APQP 时指定子零件负责人，默认 SQD。"],
        ["生成粒度", "子零件开发任务按“零件 + 供应商方案”生成，不仅按零件生成；同一零件一供、二供均需开发时生成两条独立任务。"],
        ["候选预览", "一键创建 APQP 前展示任务类型、层级、零件、供应商、供应类型、负责人、起始 Gate、是否阻塞。"],
        ["任务列表字段", "APQP 编号、层级、零件/材料名称、供应商、供应类型、负责人、来源 BOM、起始 Gate、当前 Gate、是否阻塞 Gate、状态。"],
        ["展示规则", "L1 及其子零件的 Gate 阶段任务统一展示在项目 APQP 的 Gate 阶段任务下；操作权限仍按具体 APQP 负责人控制。"],
        ["推进条件", "当前 Gate 下项目级任务、L1 APQP 任务、子零件 APQP 任务全部完成后，才可提交审核。"],
    ], [2100, 7200])

    doc.add_heading("6.9 待办任务中心", level=2)
    table(doc, ["项目", "需求说明"], [
        ["使用角色", "PM、AE、SCM、AQE、SQD、任务负责人、审核小组成员。"],
        ["角色切换", "顶部支持按角色查看待办；关键节点方案版包含 PM、AE、SCM、AQE、SQD、任务负责人、审核小组。"],
        ["统计摘要", "展示来自关键节点与 Gate 阶段的待办数量、BOM/APQP/子零件任务数量、阻塞 Gate 数量。"],
        ["列表字段", "任务名称、角色、项目/来源、阶段、触发原因、阻塞、操作；普通版还包含创建时间、要求完成时间。"],
        ["筛选", "全部、待处理、已完成；关键节点方案可按角色和阻塞状态识别处理优先级。"],
        ["处理弹窗", "按任务类型展示不同表单：流程选择、BOM 处理、Gate 准入检查、任务完成审核、Gate 审核结论、Hand Over 登记。"],
        ["待办状态", "任务责任人变更时原待办关闭或转派，新责任人生成新待办；历史待办保留状态与转派原因。"],
    ], [1800, 7500])

    doc.add_heading("6.10 子任务库管理", level=2)
    table(doc, ["功能", "说明"], [
        ["阶段维护", "支持新建阶段，维护阶段名称、排序、阶段说明、状态。"],
        ["任务维护", "支持新建任务，维护任务编码、任务名称、所属分类、任务说明、状态、关键任务标记。"],
        ["默认展示规则", "任务可配置默认展示，适用系统类型，以及不同类型系统下默认展示任务。"],
        ["系统类型", "标准流程和简化流程根据产品类型生成不同 Gate 任务；Gate4-Gate6 不应对简化流程生成无效待办。"],
        ["配置保存", "超级管理员保存配置后影响后续新项目或新阶段任务生成；已生成实例是否同步调整需客户确认。"],
    ], [2200, 7100])

    doc.add_heading("6.11 子零件开发任务详情", level=2)
    p(doc, "子零件开发任务详情页面向 SQD、AQE、PM 等角色，展示单个子零件 APQP 的阶段任务、关联 BOM、客户要求、问题、8D、团队和资料。")
    table(doc, ["区域", "需求说明"], [
        ["头部信息", "展示子零件名称、任务编号、状态；支持编辑信息、暂停/终止/删除任务（按权限）。"],
        ["项目记录", "展示当前提醒、子零件状态摘要、子零件动态，并按阶段/BOM/任务/审批/供应商/状态过滤。"],
        ["项目开发任务", "按 Gate1-Gate6 展示子零件任务，支持任务填报、审核记录、任务管理、提交。"],
        ["关联生产 BOM", "展示该零件关联 BOM 信息、阶段版本、状态，支持查询、查看明细、查看记录、查看子零件 APQP 信息。"],
        ["问题/8D/资料", "与项目详情一致，但范围聚焦子零件。"],
        ["成员选择", "支持选择交接团队、变更委员会团队及成员，支持展开/收起、全选、清空、确认。"],
    ], [2100, 7200])

    doc.add_heading("6.12 项目状态与 Hand Over", level=2)
    table(doc, ["功能", "说明"], [
        ["项目暂停/终止", "项目详情支持项目状态操作，危险操作需二次确认并填写原因；状态变化写入项目动态。"],
        ["带条件通过", "Gate 审核可存在带条件通过状态；需记录关闭条件、责任人、期限和验证人，是否允许进入下一 Gate 需客户确认。"],
        ["Hand Over 登记", "最终 Gate 完成后由 PM 触发，登记交接团队、交接成员、交接内容、附件和未关闭事项。"],
        ["结束规则", "交接完成后项目进入结束节点；简化流程 Gate3 完成后是否直接进入 Hand Over 需确认。"],
    ], [2200, 7100])

    doc.add_heading("6.13 数据统计看板", level=2)
    table(doc, ["区域", "需求说明"], [
        ["筛选条件", "统计周期：2026 Q2、2026 上半年、2026 全年；流程类型：全部、标准、简化；产品系统：全部、VT、TD、PP/Chain 子零件。"],
        ["KPI", "项目总数、进行中、风险/逾期、Hand Over 完成等。"],
        ["图表", "阶段分布与门径漏斗、管理层关注摘要、OTS/PPAP/SOP 节点健康度、流程类型与风险构成、近 6 个月项目创建趋势、跨职能任务负载。"],
        ["重点项目节点清单", "项目、客户、产品系统、流程、当前阶段、客户 OTS、客户 PPAP、客户 SOP、责任 PM、状态。"],
    ], [2000, 7300])

    doc.add_heading("7. Gate Checklist 与任务模板", level=1)
    p(doc, "Gate Checklist 已按当前 UI 阶段拆分为任务模板数据。Gate 0 已合并到 CKO，重复项已去除，UI-ready 数据总计 333 条。")
    rows = get_gate_summary_rows()
    if rows:
        table(doc, ["UI 阶段", "任务数量"], rows, [2500, 2500])
    table(doc, ["任务模板字段", "说明"], [
        ["taskId", "任务模板唯一标识。"],
        ["K-PAC ID / Title", "来自 Gate Checklist 的任务编号和标题。"],
        ["function", "任务所属功能分类，如 Customer、Plan and Define、Quality Plan 等。"],
        ["uiStage", "系统 UI 阶段：CKO、PKO、Gate1-Gate6。"],
        ["requiredType", "required/optional，对应必做/可选。"],
        ["taskDescription", "阶段任务说明，若 Excel 有 Gate 特定说明则取对应阶段描述。"],
        ["documentOrEvidence", "任务交付物或证据要求。"],
        ["responsible", "默认责任部门或角色，用于任务分配范围，不直接等同可操作人。"],
    ], [2200, 7100])

    doc.add_heading("8. 状态、数据对象与字段建议", level=1)
    table(doc, ["对象", "关键字段"], [
        ["项目", "projectId、projectNo、projectName、customer、system、product、flowType、pmUserId、currentStage、projectStatus、createdBy、createdAt。"],
        ["项目成员", "projectId、userId、roleCode、departmentCode、memberStatus、joinedAt、removedAt、removedBy。"],
        ["任务实例", "taskInstanceId、projectId、apqpId、stage、templateTaskId、responsibleDept、ownerUserId、collaboratorUserIds、requiredType、status、plannedStartAt、plannedEndAt、actualCompletedAt。"],
        ["BOM 版本", "bomVersionId、projectId、version、stage、status、workflowStatus、controlCode、lockedAt、submittedBy、submittedAt。"],
        ["BOM 行", "bomNodeId、bomVersionId、parentNodeId、level、partName、qty、sharedPart、primary/secondary supply plan、apqpPolicy、blockGate、startGate、changeStatus、apqpStatus。"],
        ["供应方案", "supplyPlanId、bomNodeId、supplierId、supplierType、supplyStrategy、needDevelopment、apqpPolicy、blockGate、startGate、ownerUserId、linkedApqpId。"],
        ["APQP 实例", "apqpId、projectId、parentApqpId、bomNodeId、level、supplierId、supplyType、ownerUserId、sourceBomVersion、startGate、currentGate、blockGate、status。"],
        ["审批记录", "approvalId、projectId、stage、approvalType、result、comment、approverUserId、approvedAt、conditionItems。"],
        ["操作日志", "actorUserId、businessOwnerUserId、operationType、operationReason、beforeValue、afterValue、operatedAt。"],
    ], [1700, 7600])

    doc.add_heading("9. 业务规则汇总", level=1)
    rules = [
        "项目创建后进入 CKO；销售完成 CKO 任务后提交至 PKO。",
        "PKO 阶段必须维护客户 OTS/PPAP/SOP 与 BW OTS/PPAP/SOP 节点，并选择标准流程或简化流程。",
        "标准流程覆盖 Gate1-Gate6；简化流程只覆盖 Gate1-Gate3。",
        "BOM 未经 SCM 确认锁定前，不得生成正式 L1 APQP 和子零件 APQP。",
        "L1 APQP 必须创建；L1 负责人默认 AQE；子零件 APQP 负责人默认 SQD。",
        "子零件 APQP 按“零件 + 供应商方案”生成；一供和二供同步开发时生成独立任务。",
        "二供备案默认不生成 APQP、不阻塞 Gate；二供后续导入启动前不阻塞，启动后按阻塞型任务处理。",
        "当前 Gate 提交审核前必须校验项目级任务、L1 APQP 任务、子零件 APQP 任务、BOM 变更均满足条件。",
        "人员从项目中删除前必须校验其是否为 PM、APQP 负责人、任务责任人或审核人；存在未完成任务时需先转派。",
        "所有关键操作必须记录审计日志，包括任务提交、驳回、审核、改派、成员删除、BOM 锁定、APQP 创建、项目暂停/终止。",
        "所有用户可见弹窗不得使用浏览器默认 alert/confirm/prompt，必须使用统一弹窗组件。",
    ]
    for item in rules:
        bullet(doc, item)

    doc.add_heading("10. 非功能与交互要求", level=1)
    table(doc, ["类别", "要求"], [
        ["权限安全", "前后端均需校验；后端为最终权限判断；按钮隐藏仅作为体验优化。"],
        ["审计追溯", "BOM 版本、BOM 行变更、供应方案变化、任务生成、关闭、继承、审批判断均需留痕。"],
        ["统一弹窗", "确认、删除、驳回、业务阻塞、表单校验均使用统一弹窗，不使用浏览器默认弹窗。"],
        ["数据保留", "历史 BOM、历史待办、附件、审核记录、离项人员历史记录不得物理删除。"],
        ["易用性", "列表支持筛选、分页、固定列；长表格需横向滚动；状态以标签区分。"],
        ["性能", "项目列表、BOM 大表、任务列表需支持分页或虚拟滚动；筛选条件应服务端分页查询。"],
        ["导入导出", "BOM 支持受控导出；问题清单支持批量导入/导出；后续模板格式需确认。"],
        ["提醒", "待办、逾期、BOM 变更阻塞、Gate 审核、关键节点临近应支持消息提醒，具体渠道待确认。"],
    ], [1700, 7600])

    doc.add_heading("11. 验收标准", level=1)
    table(doc, ["模块", "验收标准"], [
        ["项目创建", "必填校验正确；项目创建后可进入 CKO；项目列表和详情展示一致。"],
        ["流程流转", "标准/简化流程阶段生成正确；非法跳 Gate 被阻止；阶段审核结果正确影响项目状态。"],
        ["任务管理", "Gate Checklist 模板能生成实例；责任人、协办人、审核人权限正确；待办状态随任务流转同步。"],
        ["BOM 管理", "AE 可编辑并提交；SCM 可确认并锁定；版本号、历史版本、受控编码、变更影响正确。"],
        ["APQP", "锁定 BOM 后可预览并生成 APQP；按零件+供应方案生成；L1、子零件负责人和阻塞规则正确。"],
        ["Gate 准入", "存在未完成必做任务、阻塞型子零件任务、未处理 BOM 变更时不能提交审核，并能提示具体原因。"],
        ["权限", "非项目成员不可见；同部门非责任人不可操作；成员删除前校验未完成任务和负责人身份。"],
        ["看板", "阶段分布、节点健康度、风险项目、任务负载与筛选条件联动正确。"],
    ], [1800, 7500])

    doc.add_heading("12. 待确认问题清单", level=1)
    questions = [
        "简化流程 Gate1-Gate3 完成后，是否直接进入 Hand Over？",
        "带条件通过 Gate 时是否允许进入下一 Gate？条件关闭是否影响后续 Gate 审核？",
        "PM 是否允许代替任务责任人填报和提交任务？若允许是否必须填写代办原因？",
        "协办人是否可以上传附件、编辑草稿或提交任务？",
        "职能经理是否有权改派本部门任务责任人，还是只能 PM 改派？",
        "L1 APQP 负责人默认 AQE 是否可由 PM 改为其他成员？",
        "子零件 APQP 是否允许多个 SQD 协同？主责/协办权限如何划分？",
        "任务模板调整后，是否影响已生成项目任务实例？",
        "BOM 导入模板、导出模板和受控 PDF 格式是否需要客户标准模板？",
        "消息提醒渠道是否包括站内、邮件、企业微信/钉钉，以及逾期升级规则？",
        "是否需要与 ERP/PLM/MES/SAP 或供应商系统集成？若需要，请确认接口边界。",
        "客户要求、问题清单、8D 是否纳入本期开发范围，还是作为项目详情通用能力预留？",
    ]
    for item in questions:
        numbered(doc, item)

    doc.add_heading("附录 A：页面原型文件清单", level=1)
    table(doc, ["文件", "用途"], [
        ["项目开发管理.html", "项目列表与新增项目。"],
        ["index-project-record-bom-supply-apqp.html", "项目详情、BOM、供应方案、APQP、项目记录、团队、资料库。"],
        ["待办任务中心-关键节点方案.html", "按关键节点、BOM/APQP、Gate 阻塞规则生成待办入口。"],
        ["task-management.html", "阶段任务管理与子任务库选择。"],
        ["子任务库管理.html", "阶段与任务模板配置。"],
        ["子零件开发任务详情_调整版.html", "子零件 APQP/开发任务详情。"],
        ["项目数据统计看板.html", "项目阶段、节点、风险、负载统计看板。"],
    ], [3300, 6000])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("博格华纳-项目开发管理模块需求规格说明书 V1.0")
    set_run_font(r, size=9, color="666666")

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
